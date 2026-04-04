"""
DOCX (Word Document) Extraction Service
"""
import logging
from typing import Dict
from docx import Document
import base64
from io import BytesIO
from app.services.math_detector import math_detector

logger = logging.getLogger(__name__)


class DOCXExtractor:
    """Extract text from Word documents"""
    
    async def extract(self, file_path: str) -> Dict:
        """
        Extract content from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dict with text and metadata
        """
        try:
            doc = Document(file_path)
            
            # Extract all paragraphs
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract tables
            table_content = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    table_content.append(row_text)
            
            # Extract images (for potential math equations)
            images = []
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        img_base64 = base64.b64encode(image_data).decode('utf-8')
                        images.append({"data": img_base64})
                    except Exception:
                        pass
            
            full_text = "\n\n".join(text_content)
            if table_content:
                full_text += "\n\n### Tables ###\n" + "\n".join(table_content)
            
            return {
                "text": full_text,
                "images": images,
                "math_expressions": math_detector.extract_math_expressions(full_text),
                "math_text": math_detector.to_math_text(full_text),
                "metadata": {
                    "paragraph_count": len(text_content),
                    "table_count": len(doc.tables),
                    "image_count": len(images)
                }
            }
            
        except Exception as e:
            logger.error(f"Error extracting DOCX: {str(e)}")
            raise