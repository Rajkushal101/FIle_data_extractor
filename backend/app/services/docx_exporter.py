"""
DOCX Export Service
Generates formatted Word documents from Markdown content
"""

import logging
from pathlib import Path
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

logger = logging.getLogger(__name__)


class DOCXExporter:
    """Export structured notes to Word DOCX format"""
    
    def __init__(self):
        pass
    
    async def export_to_docx(
        self,
        markdown_text: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Export Markdown content to DOCX
        
        Args:
            markdown_text: Markdown-formatted text
            metadata: Document metadata (title, author, etc.)
            
        Returns:
            DOCX file as bytes
        """
        try:
            doc = Document()
            
            # Set up document metadata
            metadata = metadata or {}
            doc.core_properties.title = metadata.get('title', 'Document Notes')
            doc.core_properties.author = metadata.get('author', 'File Data Extractor')
            
            # Add title page
            self._add_title_page(doc, metadata)
            
            # Process markdown content
            self._process_markdown(doc, markdown_text)
            
            # Save to bytes
            from io import BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            logger.info(f"Successfully generated DOCX ({buffer.getbuffer().nbytes} bytes)")
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            raise
    
    def _add_title_page(self, doc: Document, metadata: Dict):
        """Add title page to document"""
        title = metadata.get('title', 'Document Notes')
        author = metadata.get('author', '')
        date = metadata.get('date', '')
        
        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        
        # Author
        if author:
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(author)
            author_run.font.size = Pt(14)
            author_run.font.italic = True
        
        # Date
        if date:
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.add_run(date)
            date_run.font.size = Pt(12)
            date_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Page break
        doc.add_page_break()
    
    def _process_markdown(self, doc: Document, markdown_text: str):
        """Process markdown content and add to document"""
        lines = markdown_text.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Headers
            if line.startswith('# '):
                self._add_heading(doc, line[2:], level=1)
            elif line.startswith('## '):
                self._add_heading(doc, line[3:], level=2)
            elif line.startswith('### '):
                self._add_heading(doc, line[4:], level=3)
            
            # Code blocks
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                self._add_code_block(doc, '\n'.join(code_lines))
            
            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                self._add_list_item(doc, line[2:])
            elif re.match(r'^\d+\.\s', line):
                match = re.match(r'^\d+\.\s(.+)', line)
                if match:
                    self._add_numbered_list_item(doc, match.group(1))
            
            # Math blocks (display)
            elif line.startswith('$$'):
                math_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('$$'):
                    math_lines.append(lines[i])
                    i += 1
                self._add_math_block(doc, '\n'.join(math_lines))
            
            # Blockquotes
            elif line.startswith('> '):
                self._add_blockquote(doc, line[2:])
            
            # Horizontal rules
            elif line == '---' or line == '***':
                self._add_horizontal_rule(doc)
            
            # Regular paragraphs
            elif line:
                self._add_paragraph(doc, line)
            
            i += 1
    
    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """Add heading to document"""
        heading = doc.add_heading(text, level=level)
        return heading
    
    def _add_paragraph(self, doc: Document, text: str):
        """Add paragraph with inline formatting"""
        para = doc.add_paragraph()
        
        # Process inline math: $...$
        parts = re.split(r'(\$[^\$]+\$)', text)
        
        for part in parts:
            if part.startswith('$') and part.endswith('$'):
                # Math expression
                math_run = para.add_run(part)
                math_run.font.italic = True
                math_run.font.color.rgb = RGBColor(0, 0, 139)
            else:
                # Regular text with bold/italic
                self._add_formatted_text(para, part)
        
        return para
    
    def _add_formatted_text(self, para, text: str):
        """Add text with bold and italic formatting"""
        # Bold: **text**
        bold_parts = re.split(r'(\*\*[^\*]+\*\*)', text)
        
        for bold_part in bold_parts:
            if bold_part.startswith('**') and bold_part.endswith('**'):
                run = para.add_run(bold_part[2:-2])
                run.font.bold = True
            else:
                # Italic: *text*
                italic_parts = re.split(r'(\*[^\*]+\*)', bold_part)
                for italic_part in italic_parts:
                    if italic_part.startswith('*') and italic_part.endswith('*'):
                        run = para.add_run(italic_part[1:-1])
                        run.font.italic = True
                    else:
                        para.add_run(italic_part)
    
    def _add_code_block(self, doc: Document, code: str):
        """Add code block with monospace font"""
        para = doc.add_paragraph()
        para.style = 'Quote'
        run = para.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        return para
    
    def _add_list_item(self, doc: Document, text: str):
        """Add bullet list item"""
        para = doc.add_paragraph(style='List Bullet')
        self._add_formatted_text(para, text)
        return para
    
    def _add_numbered_list_item(self, doc: Document, text: str):
        """Add numbered list item"""
        para = doc.add_paragraph(style='List Number')
        self._add_formatted_text(para, text)
        return para
    
    def _add_math_block(self, doc: Document, math: str):
        """Add display math block"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"$${math}$$")
        run.font.italic = True
        run.font.color.rgb = RGBColor(0, 0, 139)
        return para
    
    def _add_blockquote(self, doc: Document, text: str):
        """Add blockquote"""
        para = doc.add_paragraph(style='Quote')
        self._add_formatted_text(para, text)
        return para
    
    def _add_horizontal_rule(self, doc: Document):
        """Add horizontal rule"""
        para = doc.add_paragraph()
        para.add_run('_' * 50)
        return para


# Singleton instance
docx_exporter = DOCXExporter()


async def export_markdown_to_docx(
    markdown_text: str,
    metadata: Optional[Dict[str, Any]] = None
) -> bytes:
    """
    Convenience function for DOCX export
    
    Args:
        markdown_text: Markdown content
        metadata: Document metadata
        
    Returns:
        DOCX file as bytes
    """
    return await docx_exporter.export_to_docx(markdown_text, metadata)
